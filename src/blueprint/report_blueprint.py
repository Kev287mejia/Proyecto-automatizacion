import os
import json
import logging
from docx import Document
import openpyxl

logging.basicConfig(level=logging.INFO, format='%(levelname)s - %(message)s')

class BlueprintExtractor:
    def __init__(self, base_dir: str):
        self.base_dir = base_dir

    def extract_word(self, file_name: str, output_name: str):
        file_path = os.path.join(self.base_dir, file_name)
        output_path = os.path.join(self.base_dir, output_name)
        
        if not os.path.exists(file_path):
            logging.error(f"Archivo no encontrado: {file_path}")
            return
            
        try:
            doc = Document(file_path)
            blueprint = {
                "tipo": "word",
                "archivo_origen": file_name,
                "estructura": []
            }
            
            current_section = None
            
            for element in doc.element.body:
                if element.tag.endswith('p'): # Paragraph
                    # Re-instantiate paragraph to check style
                    for p in doc.paragraphs:
                        if p._element == element:
                            if p.style.name.startswith('Heading'):
                                # New section
                                if current_section:
                                    blueprint["estructura"].append(current_section)
                                current_section = {
                                    "seccion": p.text.strip(),
                                    "elementos": []
                                }
                            elif current_section and p.text.strip():
                                # We only add a generic 'texto' element once per text block to avoid huge blueprints
                                if "texto" not in current_section["elementos"]:
                                    current_section["elementos"].append("texto")
                            break
                elif element.tag.endswith('tbl'): # Table
                    if current_section:
                        # Count how many tables we have to give it a unique name
                        table_count = sum(1 for el in current_section["elementos"] if el.startswith("tabla"))
                        current_section["elementos"].append(f"tabla_{table_count + 1}")
                    else:
                        # Table before any heading
                        current_section = {"seccion": "Inicio", "elementos": ["tabla_1"]}
            
            if current_section:
                blueprint["estructura"].append(current_section)

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(blueprint, f, indent=4, ensure_ascii=False)
            logging.info(f"Word blueprint generado: {output_name}")
            
        except Exception as e:
            logging.error(f"Error parseando Word: {e}")


    def extract_excel(self, file_name: str, output_name: str):
        file_path = os.path.join(self.base_dir, file_name)
        output_path = os.path.join(self.base_dir, output_name)
        
        if not os.path.exists(file_path):
            logging.error(f"Archivo no encontrado: {file_path}")
            return
            
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            blueprint = {
                "tipo": "excel",
                "archivo_origen": file_name,
                "estructura": []
            }
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                elementos = []
                
                # Intentamos detectar si hay tablas formales
                if sheet.tables:
                    for table in sheet.tables.values():
                        elementos.append(f"tabla_{table.name}")
                
                # Si no hay tablas formales pero hay datos
                if not elementos and sheet.max_row > 1:
                    elementos.append("datos_tabulares")
                    
                # Chequear gráficos
                if sheet._charts:
                    for i in range(len(sheet._charts)):
                        elementos.append(f"grafico_{i+1}")

                blueprint["estructura"].append({
                    "hoja": sheet_name,
                    "elementos": elementos
                })

            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(blueprint, f, indent=4, ensure_ascii=False)
            logging.info(f"Excel blueprint generado: {output_name}")
            
        except Exception as e:
            logging.error(f"Error parseando Excel: {e}")

if __name__ == "__main__":
    base_dir = r"C:\Users\LENOVO X1 YOGA\Videos\AUTOMATIZACION\INFORMES QUE HAGO"
    extractor = BlueprintExtractor(base_dir)
    
    extractor.extract_excel("informes.xlsx", "excel_blueprint.json")
    extractor.extract_word("Informe_Analitico_SIEA.docx", "word_blueprint.json")
