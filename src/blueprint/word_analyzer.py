import os
import logging
from docx import Document

class WordAnalyzer:
    """
    WordAnalyzer es responsable de extraer las secciones, estilos y métricas 
    avanzadas (tablas, imágenes, saltos de página, etc.) de una plantilla de Word (.docx).
    No modifica nada, solo aprende del documento.
    """
    def analyze(self, file_path: str) -> dict:
        word_data = {
            "template": os.path.basename(file_path),
            "sections": [],
            "styles": {},
            "metrics": {
                "tables": 0,
                "images": 0,
                "page_breaks": 0,
                "headers_detected": False,
                "footers_detected": False,
                "lists_detected": False
            }
        }
        
        try:
            doc = Document(file_path)
            
            # --- 1. Tablas e Imágenes ---
            word_data["metrics"]["tables"] = len(doc.tables)
            word_data["metrics"]["images"] = len(doc.inline_shapes)
            
            # --- 2. Encabezados y Pies de Página ---
            # Revisamos si existe al menos un encabezado o pie con contenido real
            for section in doc.sections:
                if section.header and any(p.text.strip() for p in section.header.paragraphs):
                    word_data["metrics"]["headers_detected"] = True
                if section.footer and any(p.text.strip() for p in section.footer.paragraphs):
                    word_data["metrics"]["footers_detected"] = True
            
            # --- 3. Párrafos (Títulos, Estilos, Saltos de Página, Listas) ---
            for p in doc.paragraphs:
                # Contar saltos de página (\x0c es el caracter de salto duro)
                if '\x0c' in p.text:
                    word_data["metrics"]["page_breaks"] += p.text.count('\x0c')
                
                style_name = p.style.name if p.style else ""
                
                # Identificar Listas y Numeración
                # Buscamos estilos que usualmente son listas, o revisamos a nivel XML
                if style_name.startswith('List') or bool(p._element.xpath('./w:pPr/w:numPr')):
                    word_data["metrics"]["lists_detected"] = True
                
                # Identificar si es un Heading para armar la estructura
                if style_name.startswith('Heading'):
                    text = p.text.strip()
                    if text and text not in word_data["sections"]:
                        word_data["sections"].append(text)
                
                # Mapear todos los estilos que encontramos
                if style_name:
                    style_key = style_name.replace(" ", "")
                    if style_key not in word_data["styles"]:
                        # Fallbacks
                        font_name = p.style.font.name if getattr(p.style, 'font', None) and p.style.font.name else "Calibri"
                        font_size_pt = p.style.font.size.pt if getattr(p.style, 'font', None) and p.style.font.size else 11
                        
                        word_data["styles"][style_key] = {
                            "font": font_name,
                            "size": int(font_size_pt) if font_size_pt else 11
                        }
                        
            return word_data
            
        except Exception as e:
            logging.error(f"Error al analizar el archivo Word {file_path}: {e}")
            raise
