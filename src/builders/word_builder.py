import os
import json
import logging
from datetime import datetime
from docx import Document

logging.basicConfig(level=logging.INFO, format='%(levelname)s - %(message)s')

class WordReportBuilder:
    def __init__(self, base_dir: str):
        self.base_dir = base_dir

    def _load_json(self, file_name: str):
        path = os.path.join(self.base_dir, file_name)
        try:
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logging.error(f"Error cargando JSON {file_name}: {e}")
            return None

    def _extract_flat_data(self, assembled_report: dict) -> dict:
        """
        Extrae todos los textos, insights y métricas del reporte ensamblado 
        y los aplana en un diccionario clave-valor para facilitar el reemplazo.
        """
        flat_data = {
            "fecha_reporte": datetime.now().strftime("%d/%m/%Y")
        }
        
        for section in assembled_report.get("sections", []):
            for item in section.get("content", []):
                # Usaremos la etiqueta (label) limpiada como marcador
                label = item.get("label", "sin_etiqueta").lower().replace(" ", "_")
                if item.get("type") in ["metric", "insight"]:
                    flat_data[label] = str(item.get("data", ""))
        
        return flat_data

    def build_report(self, template_name: str, assembled_data_name: str, output_name: str):
        template_path = os.path.join(self.base_dir, template_name)
        output_path = os.path.join(self.base_dir, output_name)
        
        # 1. Cargar datos
        assembled_data = self._load_json(assembled_data_name)
        if not assembled_data:
            return
            
        context = self._extract_flat_data(assembled_data)
        logging.info(f"Marcadores disponibles para inyección: {list(context.keys())}")
        
        # 2. Cargar Plantilla
        try:
            doc = Document(template_path)
        except Exception as e:
            logging.error(f"No se pudo cargar la plantilla Word: {e}")
            return

        # Demo: Inyectamos un marcador falso al final para demostrar que el reemplazo funciona,
        # dado que la plantilla actual probablemente aún no tiene marcadores escritos por el usuario.
        doc.add_paragraph("Demo Automática de Inyección: El reporte se generó el {{fecha_reporte}} con un total de {{total_de_registros}} registros, mostrando un {{crecimiento_de_registros_vs_mes_anterior}}.")

        # 3. Lógica de Reemplazo (preservando formato)
        # Iteramos sobre todos los párrafos y runs (los runs mantienen el formato negrita/cursiva/fuente)
        reemplazos_hechos = 0
        for paragraph in doc.paragraphs:
            if '{{' in paragraph.text and '}}' in paragraph.text:
                # El reemplazo a nivel de run es complejo porque un marcador puede estar dividido en varios runs.
                # Para simplificar en esta versión, reemplazaremos el texto del párrafo si contiene el marcador.
                # En un entorno de producción, usamos librerías como docxtpl para mantener el formato intrínseco.
                for key, value in context.items():
                    marker = f"{{{{{key}}}}}"
                    if marker in paragraph.text:
                        paragraph.text = paragraph.text.replace(marker, value)
                        reemplazos_hechos += 1
        
        # 4. Guardar archivo final
        try:
            doc.save(output_path)
            logging.info(f"Se inyectaron {reemplazos_hechos} marcadores en el texto.")
            logging.info(f"¡Reporte Word construido exitosamente! Guardado como: {output_name}")
        except Exception as e:
            logging.error(f"Error guardando el documento de Word: {e}")

if __name__ == "__main__":
    base_dir = r"C:\Users\LENOVO X1 YOGA\Videos\AUTOMATIZACION\INFORMES QUE HAGO"
    
    builder = WordReportBuilder(base_dir)
    builder.build_report(
        template_name="Plantilla_SIEA.docx",
        assembled_data_name="assembled_report.json",
        output_name="Informe_Final_SIEA.docx"
    )
